from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


WORKSPACE = Path(__file__).resolve().parents[2]
SOURCE = Path(os.environ["REPORT_DOC"])
OUTPUT = WORKSPACE / "570报告_修订版.docx"
ASSET_DIR = WORKSPACE / "docs" / "report" / "assets"
DIAGRAM = ASSET_DIR / "串级PID_前馈_梯形速度规划框图.png"


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str = "#334155",
    width: int = 5,
    dashed: bool = False,
) -> None:
    x1, y1 = start
    x2, y2 = end
    if dashed:
        segments = 14
        for index in range(segments):
            if index % 2 == 0:
                t1 = index / segments
                t2 = min((index + 1) / segments, 0.92)
                draw.line(
                    (x1 + (x2 - x1) * t1, y1 + (y2 - y1) * t1,
                     x1 + (x2 - x1) * t2, y1 + (y2 - y1) * t2),
                    fill=color,
                    width=width,
                )
    else:
        draw.line((x1, y1, x2, y2), fill=color, width=width)
    angle = __import__("math").atan2(y2 - y1, x2 - x1)
    length = 18
    spread = 0.55
    points = [
        (x2, y2),
        (x2 - length * __import__("math").cos(angle - spread),
         y2 - length * __import__("math").sin(angle - spread)),
        (x2 - length * __import__("math").cos(angle + spread),
         y2 - length * __import__("math").sin(angle + spread)),
    ]
    draw.polygon(points, fill=color)


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    lines: list[str],
    fill: str,
    outline: str,
    font: ImageFont.FreeTypeFont,
    title_font: ImageFont.FreeTypeFont | None = None,
) -> None:
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=4)
    x1, y1, x2, y2 = xy
    line_gap = 9
    heights = []
    for index, line in enumerate(lines):
        selected_font = title_font if index == 0 and title_font else font
        bbox = draw.textbbox((0, 0), line, font=selected_font)
        heights.append(bbox[3] - bbox[1])
    total_height = sum(heights) + line_gap * (len(lines) - 1)
    y = y1 + (y2 - y1 - total_height) / 2
    for index, line in enumerate(lines):
        selected_font = title_font if index == 0 and title_font else font
        bbox = draw.textbbox((0, 0), line, font=selected_font)
        width = bbox[2] - bbox[0]
        draw.text(((x1 + x2 - width) / 2, y), line, fill="#0f172a", font=selected_font)
        y += heights[index] + line_gap


def create_control_diagram(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1900, 1040), "white")
    draw = ImageDraw.Draw(image)
    font = load_font(31)
    small = load_font(27)
    bold = load_font(34, bold=True)
    label = load_font(25)

    draw.text((58, 35), "钢球串级闭环与底盘加速度前馈协同控制", fill="#0f172a", font=load_font(42, True))

    boxes = {
        "target": (55, 160, 250, 280),
        "position": (310, 145, 560, 295),
        "velocity_ref": (620, 145, 880, 295),
        "velocity": (940, 145, 1190, 295),
        "sum": (1260, 170, 1370, 270),
        "actuator": (1435, 130, 1690, 310),
        "plant": (1435, 420, 1690, 590),
        "sensor": (940, 440, 1190, 575),
        "speed_cmd": (55, 735, 260, 865),
        "profile": (325, 700, 625, 900),
        "wheel": (715, 720, 1000, 880),
        "feedforward": (1080, 720, 1370, 880),
    }

    draw_box(draw, boxes["target"], ["目标位置", "xᵣ"], "#e0f2fe", "#0284c7", font, bold)
    draw_box(draw, boxes["position"], ["位置外环 P", "eₓ → vᵣ"], "#dbeafe", "#2563eb", font, bold)
    draw_box(draw, boxes["velocity_ref"], ["目标球速", "限幅 ±vₘₐₓ"], "#dbeafe", "#2563eb", font, bold)
    draw_box(draw, boxes["velocity"], ["速度内环 PI", "eᵥ → θfb"], "#ede9fe", "#7c3aed", font, bold)
    draw.ellipse(boxes["sum"], fill="#fff7ed", outline="#ea580c", width=4)
    draw.text((1291, 184), "+", fill="#c2410c", font=load_font(54, True))
    draw_box(draw, boxes["actuator"], ["摆杆执行层", "角度斜率限制", "步进梯形运动/软限位"], "#fef3c7", "#d97706", small, bold)
    draw_box(draw, boxes["plant"], ["钢球—摆杆对象", "位置 x、速度 v"], "#dcfce7", "#16a34a", font, bold)
    draw_box(draw, boxes["sensor"], ["K230观测", "坐标换算与速度滤波"], "#ecfeff", "#0891b2", small, bold)
    draw_box(draw, boxes["speed_cmd"], ["巡线速度指令", "vcmd"], "#f1f5f9", "#475569", small, bold)
    draw_box(draw, boxes["profile"], ["底盘梯形速度规划", "起步：+a加", "匀速：0", "制动：−a减"], "#ffedd5", "#ea580c", small, bold)
    draw_box(draw, boxes["wheel"], ["循迹差速 + 轮速PI", "输出平滑底盘运动"], "#f1f5f9", "#475569", small, bold)
    draw_box(draw, boxes["feedforward"], ["加速度前馈", "θff = Kff·ap"], "#fee2e2", "#dc2626", font, bold)

    draw_arrow(draw, (250, 220), (310, 220))
    draw_arrow(draw, (560, 220), (620, 220))
    draw_arrow(draw, (880, 220), (940, 220))
    draw_arrow(draw, (1190, 220), (1260, 220))
    draw_arrow(draw, (1370, 220), (1435, 220))
    draw_arrow(draw, (1562, 310), (1562, 420))
    draw_arrow(draw, (1435, 505), (1190, 505))
    draw_arrow(draw, (1065, 440), (1065, 345))
    draw.line((1065, 345, 435, 345, 435, 295), fill="#0891b2", width=5)
    draw_arrow(draw, (435, 345), (435, 295), color="#0891b2")
    draw.line((1065, 345, 1065, 295), fill="#0891b2", width=5)
    draw_arrow(draw, (1065, 345), (1065, 295), color="#0891b2")
    draw.text((590, 305), "反馈 x", fill="#0e7490", font=label)
    draw.text((1078, 330), "反馈 v", fill="#0e7490", font=label)

    draw_arrow(draw, (260, 800), (325, 800))
    draw_arrow(draw, (625, 800), (715, 800))
    draw_arrow(draw, (1000, 800), (1080, 800), color="#dc2626")
    draw.text((1018, 747), "ap=Δvp/Δt", fill="#b91c1c", font=label)
    draw.line((1225, 720, 1225, 350, 1315, 350, 1315, 270), fill="#dc2626", width=5)
    draw_arrow(draw, (1315, 350), (1315, 270), color="#dc2626")
    draw.text((1238, 365), "提前抵消惯性扰动", fill="#b91c1c", font=label)
    draw_arrow(draw, (1000, 765), (1435, 540), color="#64748b", width=4, dashed=True)
    draw.text((1125, 600), "车体加/减速形成扰动", fill="#475569", font=label)

    draw.text(
        (55, 965),
        "梯形规划限制扰动幅值，前馈负责提前补偿，串级闭环负责消除模型误差与剩余偏差。",
        fill="#334155",
        font=font,
    )
    image.save(path, dpi=(220, 220), optimize=True)


def set_run_font(run, name: str, size: float, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def replace_text(paragraph: Paragraph, text: str) -> Paragraph:
    clear_paragraph(paragraph)
    paragraph.add_run(text)
    return paragraph


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    inserted = Paragraph(new_element, paragraph._parent)
    if style:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def find_prefix(document: Document, prefix: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def add_image_after(paragraph: Paragraph, image_path: Path, width: float) -> Paragraph:
    image_paragraph = insert_after(paragraph)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.first_line_indent = Pt(0)
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    return image_paragraph


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.18)
    section.bottom_margin = Inches(0.79)
    section.left_margin = Inches(0.87)
    section.right_margin = Inches(0.87)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(22)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ["Report Formula", "Report Caption"]:
        if style_name not in document.styles:
            document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    formula_style = document.styles["Report Formula"]
    formula_style.font.name = "宋体"
    formula_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    formula_style.font.size = Pt(11)
    formula_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    formula_style.paragraph_format.line_spacing = Pt(20)
    formula_style.paragraph_format.space_before = Pt(2)
    formula_style.paragraph_format.space_after = Pt(2)

    caption_style = document.styles["Report Caption"]
    caption_style.font.name = "宋体"
    caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption_style.font.size = Pt(10.5)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    caption_style.paragraph_format.line_spacing = Pt(18)
    caption_style.paragraph_format.space_before = Pt(2)
    caption_style.paragraph_format.space_after = Pt(3)

    heading_pattern = re.compile(r"^\d+(?:\.\d+)?\s+")
    formula_prefixes = ("x =", "ẍ", "eₓ", "eᵥ", "vₚ", "aₚ", "θcmd")
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        paragraph.paragraph_format.widow_control = True
        if index == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            for run in paragraph.runs:
                set_run_font(run, "黑体", 18, True)
            continue
        if heading_pattern.match(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(22)
            for run in paragraph.runs:
                set_run_font(run, "黑体", 15, True)
            continue
        if text.startswith("图 ") or text.startswith("表 "):
            paragraph.style = document.styles["Report Caption"]
            paragraph.paragraph_format.first_line_indent = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, "宋体", 10.5, False)
            continue
        if text.startswith(formula_prefixes):
            paragraph.style = document.styles["Report Formula"]
            paragraph.paragraph_format.first_line_indent = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, "宋体", 11, False)
            continue
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(1)
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(22)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, "宋体", 12, False)

    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            tr_properties = row._tr.get_or_add_trPr()
            if tr_properties.find(qn("w:cantSplit")) is None:
                tr_properties.append(OxmlElement("w:cantSplit"))
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    paragraph.paragraph_format.line_spacing = Pt(17)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        set_run_font(run, "宋体", 9.5, row_index == 0)

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def revise_report() -> None:
    create_control_diagram(DIAGRAM)
    document = Document(SOURCE)
    for style_name in ["Report Formula", "Report Caption"]:
        if style_name not in document.styles:
            document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    document.core_properties.title = "车载平衡滚球运动控制系统（H题）设计报告"
    document.core_properties.subject = "串级PID、加速度前馈与梯形加减速协同控制"

    replace_text(
        document.paragraphs[1],
        "摘要：本系统面向车载平衡滚球运动控制任务，以MSPM0G3507为主控，采用六路红外传感器完成黑线循迹，K230实时测量钢球位置与速度，MS42步进电机经齿轮齿条机构调节摆杆倾角。针对小车起步和制动时惯性力引起的钢球偏移，设计位置P—速度PI串级控制：位置外环把位置误差转换为受限目标球速，速度内环输出摆杆倾角；底盘采用梯形速度规划限制加减速度，并将规划加速度作为前馈量叠加到倾角指令。梯形规划降低扰动幅值，前馈提前抵消惯性作用，串级闭环消除模型误差和剩余偏差，从而兼顾静态定点与行驶中任意位置保持。测试表明，系统能完成循迹停车、±5 cm定点及行驶保持任务。",
    )
    replace_text(
        document.paragraphs[2],
        "关键词：循迹小车；滚球平衡；串级PID；加速度前馈；梯形速度规划；K230视觉",
    )
    replace_text(
        document.paragraphs[5],
        "本设计采用模块化方案。底盘控制器读取六路红外传感器的巡线偏差，经循迹差速与双轮速度闭环驱动小车；K230通过固定在摆杆上方的摄像头输出钢球位置和速度，主控采用位置P—速度PI串级控制计算摆杆倾角，再由MS42步进电机和齿轮齿条机构执行。为解决静止调好的控制器在小车起步、匀速和制动阶段表现不一致的问题，底盘速度指令不直接阶跃，而是经过梯形加减速规划；规划器同步输出可预知的车体加速度，作为前馈量送入滚球倾角指令。OLED显示运行时间和状态，蜂鸣器用于启动、到点和故障提示。",
    )

    replace_text(
        find_prefix(document, "系统由循迹小车"),
        "系统由循迹小车、平衡摆杆、视觉检测、控制电路和图传记录五部分组成。小车在平面内完成路径跟踪，摆杆在车体上方提供钢球滚动通道，摆杆左端铰接，另一端由步进电机驱动的齿轮齿条机构升降。控制链分为底盘运动链与钢球平衡链：前者生成梯形速度轨迹并完成巡线和轮速闭环，后者根据K230观测执行位置—速度串级闭环；两条控制链通过“规划加速度”连接，使摆杆能在底盘加减速发生的同时进行前馈补偿，而不是等钢球产生明显位移后再纠偏。各模块接口集中在自制PCB上，便于调试、替换和封装。",
    )
    replace_text(
        find_prefix(document, "采用分层控制结构"),
        "滚球控制比较了单位置环PID和串级PID两种方案。单环方案结构简单，但钢球—摆杆对象近似具有双积分特性，位置微分项又容易放大视觉坐标抖动；控制增益偏小时制动恢复慢，偏大时目标点附近易往复振荡。最终采用位置P—速度PI串级结构：外环决定钢球应以多快的速度靠近目标，内环利用实测速度形成阻尼并输出倾角，积分限幅用于补偿坡度零偏和摩擦不对称。底盘侧采用梯形加减速代替速度阶跃，并把规划加速度作为前馈量叠加到倾角指令。该复合方案把“限制扰动、提前补偿、反馈纠偏”分开处理，更适合同时完成静态定点和行驶保持任务。",
    )

    replace_text(
        find_prefix(document, "K230对摄像头画面"),
        "K230对摄像头画面进行预处理，提取钢球候选区域并计算质心。摆杆中心O和两端刻度用于建立一维标定，K230把全长约250 mm的有效行程归一化为−50.00～+50.00，并以放大100倍的定点数q发送。主控以中心为原点，将坐标换算为毫米位置：",
    )
    replace_text(find_prefix(document, "x = (p"), "x = (q / 5000)·Lh")
    replace_text(
        find_prefix(document, "其中 p₀"),
        "式中Lh为摆杆半长，默认125 mm。主控只在视觉帧序号变化时更新速度：若K230同时给出有效速度则直接换算，否则按相邻新帧的位置差和真实帧间隔计算速度，并进行一阶低通滤波。这样可避免在100 Hz控制周期内对约25帧/s的重复图像反复差分，减少速度尖峰。连续丢失视觉约100 ms时，控制器退出运行并令摆杆回中，防止用陈旧位置持续施加倾角。",
    )

    replace_text(
        find_prefix(document, "将摆杆倾角记为θ"),
        "将摆杆倾角记为θ，钢球沿凹槽方向的位置和速度分别记为x、v，车体沿摆杆方向的加速度记为ac。在小角度、纯滚动近似下，随车坐标系中的钢球动力学可写为：",
    )
    replace_text(find_prefix(document, "ẍ"), "ẍ ≈ (5/7)(gθ－ac)")
    replace_text(
        find_prefix(document, "控制器以目标位置"),
        "由上式可知，位置对倾角近似为双积分对象，仅用位置误差直接控制倾角时阻尼不足。本系统采用串级PID（位置P—速度PI）。位置外环计算位置误差ex=xr−x，并将其转换为目标球速vr：",
    )
    replace_text(find_prefix(document, "θᵣ ="), "eₓ = xᵣ－x，  vᵣ = sat(Kₓeₓ，±vmax)")
    replace_text(
        find_prefix(document, "θᵣ 经限幅"),
        "速度内环计算ev=vr−v，并输出反馈倾角θfb：",
    )
    anchor = find_prefix(document, "速度内环计算")
    anchor = insert_after(anchor, "eᵥ = vᵣ－v，  θfb = Kvp·eᵥ＋Kvi∫eᵥdt", "Report Formula")
    anchor = insert_after(
        anchor,
        "外环目标速度限幅使钢球远离目标时也不会无限加速；内环速度反馈等效增加阻尼，避免直接对噪声较大的位置量作微分。速度积分设置上下限，目标改变时清零，防止饱和后的积分累积。钢球连续进入目标±5 mm范围15个控制周期后判为稳定。该结构既能在静态任务中消除位置偏差，也为行驶扰动提供快速的速度抑制。",
    )

    heading_33 = insert_after(anchor, "3.3 梯形加减速与车体加速度前馈")
    anchor = insert_after(
        heading_33,
        "若底盘速度从0直接跳变到巡航值，车轮会在短时间内产生较大加速度，钢球受到与车体运动方向相反的惯性作用；急停时扰动方向反转，单靠视觉闭环只能在钢球已经偏移后补救。为此，巡线基准速度vp采用受限斜率的梯形规划。每个控制周期使当前规划速度按允许的最大变化量逼近指令速度：",
    )
    anchor = insert_after(anchor, "vₚ(k) = approach[vₚ(k－1)，vcmd，alim·Δt]", "Report Formula")
    anchor = insert_after(
        anchor,
        "加速时alim取设定加速度a加，停车请求到来后取设定减速度a减。于是底盘经历匀加速、匀速和匀减速三个阶段，软停车直到规划速度降为0后再停止电机。规划器同时由相邻拍速度差计算加速度：",
    )
    anchor = insert_after(anchor, "aₚ(k) = [vₚ(k)－vₚ(k－1)] / Δt", "Report Formula")
    anchor = insert_after(
        anchor,
        "ap是控制器自身生成的量，早于视觉观察到钢球偏移，因此将其乘以前馈系数Kff得到补偿倾角。当前代码仅在规划速度超过小阈值后启用前馈，避免车轮尚未建立实际运动时提前推球；加速时ap为正，减速时自动变为负，补偿方向随之反转：",
    )
    anchor = insert_after(anchor, "θcmd = θfb＋θff，  θff = Kff·aₚ", "Report Formula")
    anchor = insert_after(
        anchor,
        "三种措施作用互补：梯形轨迹先限制加减速度及惯性扰动幅值；加速度前馈在起步和刹车瞬间预置反向倾角，降低钢球初始位移；串级闭环再根据实际位置、速度消除传动间隙、轮胎打滑、坡度零偏和前馈模型误差。倾角命令进入执行层后还要经过角度变化率限制、传动比与零点换算以及步进电机梯形运动和软限位，从而避免执行器丢步或撞击机械端点。完整信号流如图3所示。",
    )
    diagram_paragraph = add_image_after(anchor, DIAGRAM, 6.18)
    insert_after(diagram_paragraph, "图 3 串级PID、加速度前馈与梯形速度规划协同控制框图", "Report Caption")

    replace_text(find_prefix(document, "3.3 小车循迹控制"), "3.4 小车循迹控制")
    replace_text(
        find_prefix(document, "设第i个灰度通道"),
        "六路红外传感器按横向位置赋予−9～+9的权值，当前检测到黑线的通道权值取平均，并经一阶低通得到巡线误差el。巡线PID输出差速修正量Δv，左右轮目标速度分别为vL=vp+Δv、vR=vp−Δv，其中基准速度vp不是阶跃值，而是3.3节所述梯形规划速度。底层双轮速度PI根据编码器实测速度修正PWM，并叠加速度前馈和静摩擦补偿。停车时状态机把巡线目标速度改为0，统一减速度将vp平滑降至0；连续丢线达到阈值则停止并退出，避免在无轨迹依据时继续运动。",
    )

    replace_text(
        find_prefix(document, "摆杆按赛题要求"),
        "摆杆按赛题要求采用长度约25 cm的PPR管改造，保持凹槽内壁光滑，钢球直径约1 cm。摆杆左端采用铰接固定，另一端连接齿轮齿条升降机构。MS42步进电机安装在车体中部，通过驱动齿轮与竖直齿条啮合，将电机转动转换为摆杆端部的升降位移；导向件和安装板限制机构偏摆。龙门架从摆杆两侧与车体连接，摄像头固定在顶部横梁上，视场覆盖整个凹槽。SolidWorks建模与实际搭建的对应关系如图4所示。",
    )
    replace_text(find_prefix(document, "图 3 龙门架"), "图 4 龙门架与齿轮齿条结构的建模及实物对照")
    replace_text(find_prefix(document, "图 4 控制系统"), "图 5 控制系统原理图")
    replace_text(find_prefix(document, "图 5 控制系统"), "图 6 控制系统PCB布局图（约87.9 mm×82.2 mm）")
    replace_text(
        find_prefix(document, "软件采用定时任务"),
        "软件采用100 Hz定时更新与事件状态机结合的结构。每拍先更新K230数据和钢球位置、速度估计，再推进巡线梯形速度轨迹与轮速闭环，随后执行滚球串级控制并把本拍规划加速度送入前馈支路，最后由摆杆执行层更新步进目标。主循环负责按键、OLED、蜂鸣器、任务计时和模式切换。位置外环增益、速度PI增益、最大球速、前馈系数、前馈启用阈值以及底盘加减速度均可通过蓝牙调参接口在线修改，并可由遥测同时观察球位置、目标球速、倾角和底盘规划加速度。任何串口超时、坐标越界、连续丢线或执行异常均进入故障保护。",
    )
    replace_text(find_prefix(document, "图 6 软件"), "图 7 软件主流程图")

    replace_text(
        find_prefix(document, "测试使用赛题环形"),
        "测试使用赛题环形黑线场地、钢球、摆杆刻度、直尺、系统计时、K230图像记录装置和上位机遥测。每项测试至少重复3次，记录完成时间、最大位置偏差和是否满足指标。测试前检查车身尺寸不超过35 cm×25 cm、摆杆高度h≥5 cm、图传画面覆盖全槽，并完成钢球坐标零点、摆杆水平零点、巡线阈值和轮速标定。控制参数按“先内环、再外环、最后前馈”的顺序整定：先在静止底盘上调速度PI和位置P，再启用梯形加减速，最后根据起步与刹车方向调节Kff。",
    )
    replace_text(
        find_prefix(document, "第1项测试在环形"),
        "第1项测试检查图传实时性与录像完整性。第2项记录A点出发至回到A点的总时间和停车偏差。第3项在小车静止时按中心、+5 cm、中心、−5 cm的顺序改变目标，记录到达时间和最大误差。第4～6项分别对应A到B中心保持、整圈中心保持和整圈任意位置保持；除最终稳态误差外，单独回看起步加速段与终点制动段的最大偏移，并对照遥测中的规划速度vp、规划加速度ap、目标球速vr和倾角θcmd，确认前馈方向及减速切换正确。",
    )
    table_heading = find_prefix(document, "5.2 测试数据")
    insert_after(table_heading, "表 3 系统测试结果", "Report Caption")

    replace_text(
        find_prefix(document, "本设计围绕车载平衡滚球任务"),
        "本设计完成了循迹、视觉定位、摆杆执行和图传记录的一体化车载滚球系统。控制部分不再把滚球稳定笼统表述为单位置PID，而是按现有程序实现位置P—速度PI串级控制，并由底盘梯形速度规划提供加速度前馈：外环保证目标位置，内环提供速度阻尼，前馈针对起步和制动的惯性扰动，执行层负责角度斜率、步进梯形运动与软限位。测试结果表明，系统能够完成一圈停车、静态±5 cm定点、A到B中心保持、整圈中心保持和任意位置保持，三次测试均满足题目时间与误差要求。",
    )
    conclusion = find_prefix(document, "本设计完成了循迹")
    insert_after(
        conclusion,
        "现阶段误差主要来自视觉帧率低于控制频率、摄像头支架残余振动、齿轮齿条回差以及轮胎打滑造成的实际加速度与规划值不一致。后续可增加IMU纵向加速度与编码器速度融合，对Kff按加速、减速分段标定，并采用带前馈校正的S形速度曲线进一步减小加速度突变；同时扩大不同电量、不同场地摩擦条件下的重复测试，以提高封闭比赛环境中的鲁棒性。",
    )

    module_table = document.tables[0]
    set_cell_text(module_table.cell(1, 2), "读取传感器，执行状态机、串级闭环与前馈控制")
    set_cell_text(module_table.cell(2, 2), "检测钢球位置与速度并输出坐标帧")
    set_cell_text(module_table.cell(4, 2), "经齿轮齿条调节摆杆倾角")

    format_document(document)
    document.save(OUTPUT)
    print(OUTPUT)
    print(DIAGRAM)


if __name__ == "__main__":
    revise_report()
