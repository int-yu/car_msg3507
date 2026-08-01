from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


WORKSPACE = Path(__file__).resolve().parents[2]
SOURCE = WORKSPACE / "570报告_修订版.docx"
OUTPUT = WORKSPACE / "570报告_模板格式修订版.docx"


def set_run_font(run, east_asia: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def replace_text(paragraph: Paragraph, text: str) -> Paragraph:
    clear_paragraph(paragraph)
    paragraph.add_run(text)
    return paragraph


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    if style:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def delete_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)
    paragraph._p = paragraph._element = None


def find_prefix(document: Document, prefix: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def configure_style(style, east_asia: str, size_pt: float, bold: bool = False) -> None:
    style.font.name = "Times New Roman"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size_pt)
    style.font.bold = bold


def set_template_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    configure_style(normal, "宋体", 12, False)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(24.1)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(22)

    heading1 = document.styles["Heading 1"]
    configure_style(heading1, "黑体", 15, True)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading1.paragraph_format.first_line_indent = Pt(0)
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(4)
    heading1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading1.paragraph_format.line_spacing = Pt(22)
    heading1.paragraph_format.keep_with_next = True

    heading2 = document.styles["Heading 2"]
    configure_style(heading2, "黑体", 14, True)
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading2.paragraph_format.first_line_indent = Pt(0)
    heading2.paragraph_format.space_before = Pt(4)
    heading2.paragraph_format.space_after = Pt(2)
    heading2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading2.paragraph_format.line_spacing = Pt(22)
    heading2.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    configure_style(caption, "宋体", 10.5, False)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(3)
    caption.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    caption.paragraph_format.line_spacing = Pt(18)
    caption.paragraph_format.keep_with_next = True

    if "Equation" not in document.styles:
        document.styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    equation = document.styles["Equation"]
    configure_style(equation, "宋体", 12, False)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    equation.paragraph_format.first_line_indent = Pt(0)
    equation.paragraph_format.space_before = Pt(4)
    equation.paragraph_format.space_after = Pt(4)
    equation.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    equation.paragraph_format.keep_together = True


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.1)
    section.different_first_page_header_footer = False

    header = section.header
    for paragraph in list(header.paragraphs)[1:]:
        delete_paragraph(paragraph)
    clear_paragraph(header.paragraphs[0])

    footer = section.footer
    for paragraph in list(footer.paragraphs)[1:]:
        delete_paragraph(paragraph)
    paragraph = footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    set_run_font(run, "宋体", 9, False)
    run._r.append(begin)
    run._r.append(instruction)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    cached = OxmlElement("w:t")
    cached.text = "1"
    run._r.append(cached)
    run._r.append(end)


def format_front_page(document: Document) -> None:
    title = document.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    title.paragraph_format.line_spacing = Pt(26)
    for run in title.runs:
        set_run_font(run, "黑体", 18, True)

    abstract = document.paragraphs[1]
    abstract_text = abstract.text.strip()
    if abstract_text.startswith("摘要："):
        abstract_text = abstract_text[3:]
    abstract_heading = abstract.insert_paragraph_before("设计报告摘要")
    abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abstract_heading.paragraph_format.first_line_indent = Pt(0)
    abstract_heading.paragraph_format.space_before = Pt(4)
    abstract_heading.paragraph_format.space_after = Pt(12)
    for run in abstract_heading.runs:
        set_run_font(run, "黑体", 15, True)
    replace_text(abstract, abstract_text)
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.first_line_indent = Pt(24.1)
    for run in abstract.runs:
        set_run_font(run, "宋体", 12, False)

    keywords = document.paragraphs[3]
    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    keywords.paragraph_format.first_line_indent = Pt(0)
    for run in keywords.runs:
        set_run_font(run, "宋体", 12, False)
    page_break = insert_after(keywords)
    page_break.paragraph_format.first_line_indent = Pt(0)
    page_break.add_run().add_break(WD_BREAK.PAGE)


def restructure_chapters(document: Document) -> None:
    replace_text(find_prefix(document, "1 引言"), "一、引言").style = "Heading 1"

    overall = find_prefix(document, "2.1 总体方案")
    overall.insert_paragraph_before("二、总体方案设计与方案论证", "Heading 1")

    theory = find_prefix(document, "3.1 视觉标定")
    theory.insert_paragraph_before("三、理论分析与控制策略", "Heading 1")

    hardware = find_prefix(document, "4.1 机械结构")
    hardware.insert_paragraph_before("四、系统硬件与机械设计", "Heading 1")

    software = find_prefix(document, "4.3 软件设计")
    replace_text(software, "五、软件设计")
    software.style = "Heading 1"
    insert_after(software, "5.1 主程序与状态机", "Heading 2")

    tests = find_prefix(document, "5 测试方案")
    replace_text(tests, "六、测试方案与测试结果")
    tests.style = "Heading 1"
    replace_text(find_prefix(document, "5.1 测试设备"), "6.1 测试环境、设备与方法")
    replace_text(find_prefix(document, "5.2 测试数据"), "6.2 测试结果")

    conclusion = find_prefix(document, "6 结论")
    original_improvement = find_prefix(document, "现阶段误差主要来自")
    improvement_text = original_improvement.text
    improvement = conclusion.insert_paragraph_before(
        "七、设计实现、调试与改进", "Heading 1"
    )
    tuning = insert_after(improvement, "7.1 参数整定与稳定性改进", "Heading 2")
    tuning = insert_after(
        tuning,
        "参数整定遵循“先执行层、再速度内环、后位置外环、最后加速度前馈”的顺序。首先校准摆杆水平零点、传动方向和齿轮比，并保证步进运动不丢步；随后在底盘静止条件下调节速度PI，使钢球速度响应具有足够阻尼，再逐步提高位置外环比例增益并设置最大目标球速。静态定位稳定后启用底盘梯形加减速，最后根据起步和制动阶段钢球偏移的方向与峰值调整前馈系数及启用阈值。该顺序避免多个环节同时变化造成参数耦合。",
    )
    reliability = insert_after(tuning, "7.2 误差来源与可靠性改进", "Heading 2")
    insert_after(reliability, improvement_text)
    delete_paragraph(original_improvement)
    replace_text(conclusion, "八、结论")
    conclusion.style = "Heading 1"

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text and text[0].isdigit() and "." in text[:4]:
            paragraph.style = "Heading 2"


def normalize_captions(document: Document) -> None:
    replacements = {
        "图 1 ": "图1  ",
        "图 2 ": "图2  ",
        "图 3 ": "图3  ",
        "图 4 ": "图4  ",
        "图 5 ": "图5  ",
        "图 6 ": "图6  ",
        "图 7 ": "图7  ",
        "表 1 ": "表1  ",
        "表 2 ": "表2  ",
        "表 3 ": "表3  ",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        for old, new in replacements.items():
            if text.startswith(old):
                replace_text(paragraph, text.replace(old, new, 1))
                text = paragraph.text.strip()
                break
        if text.startswith(("图", "表")) and len(text) > 1 and text[1].isdigit():
            paragraph.style = "Caption"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)

    captions = [
        "图1  系统总体结构框图",
        "图2  两种摆杆升降机构建模对比",
        "图3  串级PID、加速度前馈与梯形速度规划协同控制框图",
        "图4  龙门架与齿轮齿条结构的建模及实物对照",
        "图5  控制系统原理图",
        "图6  控制系统PCB布局图（约87.9 mm×82.2 mm）",
        "图7  软件主流程图",
    ]
    drawing_paragraphs = [
        paragraph for paragraph in document.paragraphs
        if paragraph._p.xpath(".//w:drawing")
    ]
    if len(drawing_paragraphs) != len(captions):
        raise ValueError(
            f"Expected {len(captions)} drawings, found {len(drawing_paragraphs)}"
        )
    paragraphs = document.paragraphs
    paragraph_indices = {id(paragraph._p): index for index, paragraph in enumerate(paragraphs)}
    for drawing, caption_text in zip(drawing_paragraphs, captions):
        drawing_index = paragraph_indices[id(drawing._p)]
        for candidate in paragraphs[drawing_index + 1:]:
            if candidate.text.strip():
                replace_text(candidate, caption_text)
                candidate.style = "Caption"
                candidate.alignment = WD_ALIGN_PARAGRAPH.CENTER
                candidate.paragraph_format.first_line_indent = Pt(0)
                break


def math_run(text: str):
    run = OxmlElement("m:r")
    properties = OxmlElement("m:rPr")
    style = OxmlElement("m:sty")
    style.set(qn("m:val"), "p")
    properties.append(style)
    run.append(properties)
    word_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Cambria Math")
    fonts.set(qn("w:hAnsi"), "Cambria Math")
    word_properties.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    word_properties.append(size)
    run.append(word_properties)
    value = OxmlElement("m:t")
    value.text = text
    run.append(value)
    return run


def math_subscript(base: str, subscript: str):
    element = OxmlElement("m:sSub")
    properties = OxmlElement("m:sSubPr")
    element.append(properties)
    base_element = OxmlElement("m:e")
    base_element.append(math_run(base))
    sub_element = OxmlElement("m:sub")
    sub_element.append(math_run(subscript))
    element.append(base_element)
    element.append(sub_element)
    return element


def math_fraction(numerator: list, denominator: list):
    element = OxmlElement("m:f")
    properties = OxmlElement("m:fPr")
    element.append(properties)
    numerator_element = OxmlElement("m:num")
    denominator_element = OxmlElement("m:den")
    for item in numerator:
        numerator_element.append(item)
    for item in denominator:
        denominator_element.append(item)
    element.append(numerator_element)
    element.append(denominator_element)
    return element


def math_accent(base: str, character: str):
    element = OxmlElement("m:acc")
    properties = OxmlElement("m:accPr")
    accent_character = OxmlElement("m:chr")
    accent_character.set(qn("m:val"), character)
    properties.append(accent_character)
    base_element = OxmlElement("m:e")
    base_element.append(math_run(base))
    element.append(properties)
    element.append(base_element)
    return element


def add_tab_stop(tabs, value: str, position: int) -> None:
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), value)
    tab.set(qn("w:pos"), str(position))
    tabs.append(tab)


def set_equation(paragraph: Paragraph, items: list, number: int) -> None:
    clear_paragraph(paragraph)
    paragraph.style = "Equation"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    properties = paragraph._p.get_or_add_pPr()
    tabs = properties.find(qn("w:tabs"))
    if tabs is not None:
        properties.remove(tabs)
    tabs = OxmlElement("w:tabs")
    add_tab_stop(tabs, "center", 4535)
    add_tab_stop(tabs, "right", 9071)
    properties.append(tabs)

    first_tab_run = OxmlElement("w:r")
    first_tab_run.append(OxmlElement("w:tab"))
    paragraph._p.append(first_tab_run)

    math = OxmlElement("m:oMath")
    for item in items:
        math.append(item)
    paragraph._p.append(math)

    number_run = OxmlElement("w:r")
    number_run.append(OxmlElement("w:tab"))
    number_text = OxmlElement("w:t")
    number_text.text = f"（{number}）"
    number_run.append(number_text)
    paragraph._p.append(number_run)


def mr(text: str):
    return math_run(text)


def ms(base: str, subscript: str):
    return math_subscript(base, subscript)


def mf(numerator: str, denominator: str):
    return math_fraction([mr(numerator)], [mr(denominator)])


def replace_formulas(document: Document) -> None:
    set_equation(
        find_prefix(document, "x = (q"),
        [mr("x="), mf("q", "5000"), mr("·"), ms("L", "h")],
        1,
    )
    set_equation(
        find_prefix(document, "ẍ"),
        [math_accent("x", "̈"), mr("≈"), mf("5", "7"), mr("(gθ−"), ms("a", "c"), mr(")")],
        2,
    )
    set_equation(
        find_prefix(document, "eₓ"),
        [
            ms("e", "x"), mr("="), ms("x", "r"), mr("−x;  "),
            ms("v", "r"), mr("=sat("), ms("K", "x"), ms("e", "x"), mr(",±"),
            ms("v", "max"), mr(")"),
        ],
        3,
    )
    set_equation(
        find_prefix(document, "eᵥ"),
        [
            ms("e", "v"), mr("="), ms("v", "r"), mr("−v;  "), ms("θ", "fb"), mr("="),
            ms("K", "vp"), ms("e", "v"), mr("+"), ms("K", "vi"), mr("∫"), ms("e", "v"), mr("dt"),
        ],
        4,
    )
    set_equation(
        find_prefix(document, "vₚ(k)"),
        [
            ms("v", "p"), mr("(k)=approach["), ms("v", "p"), mr("(k−1),"),
            ms("v", "cmd"), mr(","), ms("a", "lim"), mr("Δt]"),
        ],
        5,
    )
    numerator = [ms("v", "p"), mr("(k)−"), ms("v", "p"), mr("(k−1)")]
    denominator = [mr("Δt")]
    set_equation(
        find_prefix(document, "aₚ(k)"),
        [ms("a", "p"), mr("(k)="), math_fraction(numerator, denominator)],
        6,
    )
    set_equation(
        find_prefix(document, "θcmd"),
        [
            ms("θ", "cmd"), mr("="), ms("θ", "fb"), mr("+"), ms("θ", "ff"), mr(";  "),
            ms("θ", "ff"), mr("="), ms("K", "ff"), ms("a", "p"),
        ],
        7,
    )


def set_border(container, edge: str, value: str, size: int = 0, color: str = "000000") -> None:
    tag = qn(f"w:{edge}")
    border = container.find(tag)
    if border is None:
        border = OxmlElement(f"w:{edge}")
        container.append(border)
    border.set(qn("w:val"), value)
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int]) -> None:
    if len(widths) != len(table.columns):
        raise ValueError("Column width count does not match table")
    total_width = sum(widths)
    properties = table._tbl.tblPr
    width_element = properties.find(qn("w:tblW"))
    if width_element is None:
        width_element = OxmlElement("w:tblW")
        properties.insert(0, width_element)
    width_element.set(qn("w:type"), "dxa")
    width_element.set(qn("w:w"), str(total_width))
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "0")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().insert(0, cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(width))


def make_three_line_table(table, widths: list[int]) -> None:
    try:
        table.style = "Normal Table"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, widths)

    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    set_border(borders, "top", "single", 18)
    set_border(borders, "bottom", "single", 18)
    for edge in ["left", "right", "insideH", "insideV"]:
        set_border(borders, edge, "nil", 0)

    for row_index, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        if row_properties.find(qn("w:cantSplit")) is None:
            row_properties.append(OxmlElement("w:cantSplit"))
        if row_index == 0 and row_properties.find(qn("w:tblHeader")) is None:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            row_properties.append(repeat)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            shading = cell_properties.find(qn("w:shd"))
            if shading is None:
                shading = OxmlElement("w:shd")
                cell_properties.append(shading)
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), "E8EEF4" if row_index == 0 else "FFFFFF")
            cell_margins = cell_properties.find(qn("w:tcMar"))
            if cell_margins is None:
                cell_margins = OxmlElement("w:tcMar")
                cell_properties.append(cell_margins)
            for edge, value in [("top", 100), ("left", 120), ("bottom", 100), ("right", 120)]:
                margin = cell_margins.find(qn(f"w:{edge}"))
                if margin is None:
                    margin = OxmlElement(f"w:{edge}")
                    cell_margins.append(margin)
                margin.set(qn("w:w"), str(value))
                margin.set(qn("w:type"), "dxa")
            cell_borders = cell_properties.find(qn("w:tcBorders"))
            if cell_borders is None:
                cell_borders = OxmlElement("w:tcBorders")
                cell_properties.append(cell_borders)
            for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                set_border(cell_borders, edge, "nil", 0)
            if row_index == 0:
                set_border(cell_borders, "bottom", "single", 8)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph.paragraph_format.line_spacing = Pt(18)
                for run in paragraph.runs:
                    set_run_font(run, "宋体", 10.5, row_index == 0)


def format_body(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        paragraph.paragraph_format.widow_control = True
        if text == "车载平衡滚球运动控制系统（H题）":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            continue
        if text == "设计报告摘要":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            continue
        if paragraph.style.name in ["Heading 1", "Heading 2", "Caption", "Equation"]:
            continue
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(2)
            continue
        if not text:
            paragraph.paragraph_format.first_line_indent = Pt(0)
            continue
        if text.startswith("关键词："):
            paragraph.paragraph_format.first_line_indent = Pt(0)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Pt(24.1)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(22)
        for run in paragraph.runs:
            set_run_font(run, "宋体", 12, False)


def set_image_alt_text(document: Document) -> None:
    descriptions = [
        "系统总体结构框图",
        "两种摆杆升降机构建模对比",
        "串级PID、加速度前馈与梯形速度规划协同控制框图",
        "龙门架与齿轮齿条结构的建模及实物对照",
        "控制系统原理图",
        "控制系统PCB布局图",
        "软件主流程图",
    ]
    inline_shapes = document.inline_shapes
    if len(inline_shapes) != len(descriptions):
        raise ValueError(
            f"Expected {len(descriptions)} inline shapes, found {len(inline_shapes)}"
        )
    for shape, description in zip(inline_shapes, descriptions):
        doc_pr = shape._inline.docPr
        doc_pr.set("descr", description)
        doc_pr.set("title", description)


def finalize_settings(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)


def main() -> None:
    document = Document(SOURCE)
    set_template_styles(document)
    set_page_layout(document)
    format_front_page(document)
    restructure_chapters(document)
    normalize_captions(document)
    replace_formulas(document)
    table_widths = [
        [1500, 2550, 5021],
        [1450, 1800, 4200, 1621],
        [1350, 1750, 900, 900, 900, 3271],
    ]
    if len(document.tables) != len(table_widths):
        raise ValueError(
            f"Expected {len(table_widths)} tables, found {len(document.tables)}"
        )
    for table, widths in zip(document.tables, table_widths):
        make_three_line_table(table, widths)
    format_body(document)
    set_image_alt_text(document)
    finalize_settings(document)
    document.core_properties.title = "车载平衡滚球运动控制系统（H题）设计报告"
    document.core_properties.subject = "按竞赛模板排版的设计报告"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
